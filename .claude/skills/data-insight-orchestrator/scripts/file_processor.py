#!/usr/bin/env python3
"""文件处理器 - 支持多格式数据文件的读取"""

import os
import json
from pathlib import Path
from datetime import datetime
import pandas as pd
import openpyxl
import PyPDF2
from docx import Document


class FileProcessor:
    SUPPORTED_EXTENSIONS = {'.xlsx': 'Excel', '.xls': 'Excel', '.pdf': 'PDF', '.docx': 'Word', '.csv': 'CSV', '.txt': 'TXT'}

    def __init__(self, upload_dir='uploads'):
        self.upload_dir = Path(upload_dir)

    def scan_files(self):
        files = []
        if not self.upload_dir.exists():
            return files
        for file_path in self.upload_dir.iterdir():
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.SUPPORTED_EXTENSIONS:
                    files.append({
                        'name': file_path.name,
                        'path': str(file_path),
                        'type': self.SUPPORTED_EXTENSIONS[ext],
                        'size': self._format_size(file_path.stat().st_size),
                        'extension': ext
                    })
        return files

    def _format_size(self, size):
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} TB"

    def read_excel(self, file_path):
        result = {'type': 'Excel', 'sheets': {}, 'summary': ''}
        try:
            xl = pd.ExcelFile(file_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name)
                result['sheets'][sheet_name] = {
                    'rows': len(df),
                    'columns': list(df.columns),
                    'data': df.to_dict('records')[:100]
                }
            result['summary'] = f"{len(xl.sheet_names)} 个工作表"
        except Exception as e:
            result['error'] = str(e)
        return result

    def read_csv(self, file_path):
        result = {'type': 'CSV', 'data': [], 'summary': ''}
        try:
            df = pd.read_csv(file_path)
            result['rows'] = len(df)
            result['columns'] = list(df.columns)
            result['data'] = df.to_dict('records')[:100]
            result['summary'] = f"{len(df)} 行, {len(df.columns)} 列"
        except Exception as e:
            result['error'] = str(e)
        return result

    def read_pdf(self, file_path):
        result = {'type': 'PDF', 'pages': [], 'summary': ''}
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    result['pages'].append({'page': i + 1, 'text': text[:2000] if text else ''})
            result['summary'] = f"{len(result['pages'])} 页"
        except Exception as e:
            result['error'] = str(e)
        return result

    def read_docx(self, file_path):
        result = {'type': 'Word', 'paragraphs': [], 'tables': [], 'summary': ''}
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    result['paragraphs'].append(para.text)
            for table in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                result['tables'].append(table_data)
            result['summary'] = f"{len(result['paragraphs'])} 段落, {len(result['tables'])} 个表格"
        except Exception as e:
            result['error'] = str(e)
        return result

    def read_txt(self, file_path):
        result = {'type': 'TXT', 'content': '', 'summary': ''}
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            result['content'] = content[:10000]
            result['lines'] = len(content.split('\n'))
            result['summary'] = f"{result['lines']} 行"
        except Exception as e:
            result['error'] = str(e)
        return result

    def process_file(self, file_path):
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext in ['.xlsx', '.xls']:
            return self.read_excel(file_path)
        elif ext == '.csv':
            return self.read_csv(file_path)
        elif ext == '.pdf':
            return self.read_pdf(file_path)
        elif ext == '.docx':
            return self.read_docx(file_path)
        elif ext == '.txt':
            return self.read_txt(file_path)
        return {'error': f'不支持的文件格式: {ext}'}


def main():
    processor = FileProcessor()
    files = processor.scan_files()
    print(f"发现 {len(files)} 个文件")
    for f in files:
        print(f"  - {f['name']} ({f['type']}) - {f['size']}")


if __name__ == '__main__':
    main()